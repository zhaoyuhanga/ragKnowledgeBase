# -*- coding: utf-8 -*-
"""
Word文档解析器

本模块实现Word文档（.docx）的结构化解析。
支持提取标题、段落、表格、图片、页眉页脚等元素。
"""

from typing import Any, Dict, List, Optional

from app.parsers.base import (
    BaseParser,
    DocumentElementModel,
    ElementType,
    ImageDescription,
    TableStructure,
)


class WordParser(BaseParser):
    """
    Word文档解析器

    解析Word文档（.docx/.doc），提取结构化元素。
    """

    def __init__(self):
        """初始化Word解析器"""
        super().__init__()
        self.supported_extensions = ["docx", "doc"]

    def parse(
        self,
        file_path: str,
        version_id: int,
        document_id: int
    ) -> List[DocumentElementModel]:
        """
        解析Word文档

        Args:
            file_path: 文件路径
            version_id: 版本ID
            document_id: 文档ID

        Returns:
            解析后的元素列表
        """
        try:
            from docx import Document
        except ImportError:
            # python-docx未安装，返回空列表
            return []

        elements = []
        title_path: List[str] = []

        try:
            doc = Document(file_path)
            total_paragraphs = len(doc.paragraphs)
            # 估算页数（假设每100个段落约一页）
            estimated_pages = max(1, (total_paragraphs // 100) + 1)

            # 遍历所有段落
            for para_index, para in enumerate(doc.paragraphs):
                # 估算当前段落所在的页码
                page_no = min((para_index // 100) + 1, estimated_pages)
                element = self._parse_paragraph(para, document_id, version_id, title_path, page_no)
                if element:
                    elements.append(element)

                    # 更新标题路径
                    if element.element_type == ElementType.TITLE and element.title_level:
                        level = element.title_level
                        title_path = title_path[:level - 1] + [element.content[:50]]

            # 提取表格
            for index, table in enumerate(doc.tables):
                table_element = self._parse_table(
                    table, index, document_id, version_id, len(elements)
                )
                if table_element:
                    elements.append(table_element)

            # 提取图片
            for rel_id, rel in doc.part.rels.items():
                if "image" in rel.reltype:
                    image_element = self._parse_image(
                        rel, document_id, version_id, len(elements)
                    )
                    if image_element:
                        elements.append(image_element)

            # 提取页眉页脚
            for section in doc.sections:
                # 页眉
                header_element = self._parse_header_footer(
                    section.header, document_id, version_id, len(elements), ElementType.HEADER
                )
                if header_element:
                    elements.append(header_element)

                # 页脚
                footer_element = self._parse_header_footer(
                    section.footer, document_id, version_id, len(elements), ElementType.FOOTER
                )
                if footer_element:
                    elements.append(footer_element)

        except Exception as e:
            # 解析失败，记录错误
            import logging
            logger = logging.getLogger("rag.parser")
            logger.error(f"Word文档解析失败: {file_path}, 错误: {str(e)}")

        return elements

    def _parse_paragraph(
        self,
        para,
        document_id: int,
        version_id: int,
        title_path: List[str],
        page_no: int = 1
    ) -> Optional[DocumentElementModel]:
        """
        解析段落

        Args:
            para: 段落对象
            document_id: 文档ID
            version_id: 版本ID
            title_path: 当前标题路径
            page_no: 页码

        Returns:
            段落元素
        """
        text = para.text.strip()
        if not text:
            return None

        # 判断是否为标题
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading") or style_name.startswith("标题"):
            # 提取标题级别
            level_str = style_name.replace("Heading", "").replace("标题", "").strip()
            try:
                level = int(level_str) if level_str else 1
            except ValueError:
                level = 1

            return self._create_title_element(
                text=text,
                level=level,
                document_id=document_id,
                version_id=version_id,
                reading_order=0,
                title_path=" > ".join(title_path + [text[:50]]) if title_path else text[:50],
                page_no=page_no
            )
        elif style_name in ["List Paragraph", "List", "列表"]:
            return self._create_list_element(text, document_id, version_id, len(title_path), page_no)
        else:
            return self._create_paragraph_element(
                text=text,
                document_id=document_id,
                version_id=version_id,
                reading_order=0,
                title_path=" > ".join(title_path) if title_path else "",
                page_no=page_no
            )

    def _create_list_element(
        self,
        text: str,
        document_id: int,
        version_id: int,
        reading_order: int,
        page_no: int = 1
    ) -> DocumentElementModel:
        """创建列表元素"""
        return DocumentElementModel(
            element_id=self._generate_element_id(),
            document_id=document_id,
            version_id=version_id,
            page_no=page_no,
            element_type=ElementType.LIST,
            content=text,
            enhanced_content=text,
            reading_order=reading_order,
            confidence=0.9
        )

    def _parse_table(
        self,
        table,
        index: int,
        document_id: int,
        version_id: int,
        reading_order: int
    ) -> Optional[DocumentElementModel]:
        """
        解析表格

        Args:
            table: 表格对象
            index: 表格索引
            document_id: 文档ID
            version_id: 版本ID
            reading_order: 阅读顺序

        Returns:
            表格元素
        """
        try:
            # 提取表头
            headers = []
            if table.rows:
                header_row = table.rows[0]
                for cell in header_row.cells:
                    headers.append(cell.text.strip())

            # 提取数据行
            rows = []
            for row in table.rows[1:]:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                if any(row_data):  # 跳过空行
                    rows.append(row_data)

            # 检测合并单元格
            merged_cells = self._detect_merged_cells(table)

            # 创建表格结构
            table_structure = TableStructure(
                headers=[headers] if headers else [],
                rows=rows,
                merged_cells=merged_cells,
                row_count=len(rows) + (1 if headers else 0),
                col_count=len(headers) if headers else 0,
                caption=f"表格 {index + 1}"
            )

            return DocumentElementModel(
                element_id=self._generate_element_id(),
                document_id=document_id,
                version_id=version_id,
                element_type=ElementType.TABLE,
                content=table_structure.to_text(),
                reading_order=reading_order,
                table_structure=table_structure,
                confidence=0.95
            )
        except Exception:
            return None

    def _detect_merged_cells(self, table) -> List[Dict[str, Any]]:
        """检测合并单元格"""
        merged_cells = []
        # 简化实现，实际需要遍历单元格合并信息
        return merged_cells

    def _parse_image(
        self,
        rel,
        document_id: int,
        version_id: int,
        reading_order: int
    ) -> Optional[DocumentElementModel]:
        """
        解析图片

        Args:
            rel: 图片关系对象
            document_id: 文档ID
            version_id: 版本ID
            reading_order: 阅读顺序

        Returns:
            图片元素
        """
        try:
            # 获取图片数据
            image_part = rel.target_part
            image_bytes = image_part.blob

            # 简单的图片描述（后续可通过多模态模型增强）
            image_description = ImageDescription(
                description="Word文档中的图片",
                confidence=0.85
            )

            return DocumentElementModel(
                element_id=self._generate_element_id(),
                document_id=document_id,
                version_id=version_id,
                element_type=ElementType.IMAGE,
                enhanced_content=image_description.description,
                reading_order=reading_order,
                image_description=image_description,
                metadata={
                    "image_size": len(image_bytes),
                    "image_type": rel.target_ref.split(".")[-1] if "." in rel.target_ref else "unknown"
                },
                confidence=0.85
            )
        except Exception:
            return None

    def _parse_header_footer(
        self,
        header_footer,
        document_id: int,
        version_id: int,
        reading_order: int,
        element_type: ElementType
    ) -> Optional[DocumentElementModel]:
        """
        解析页眉页脚

        Args:
            header_footer: 页眉/页脚对象
            document_id: 文档ID
            version_id: 版本ID
            reading_order: 阅读顺序
            element_type: 元素类型（HEADER或FOOTER）

        Returns:
            页眉/页脚元素
        """
        try:
            text = header_footer.paragraphs[0].text.strip() if header_footer.paragraphs else ""
            if not text:
                return None

            return DocumentElementModel(
                element_id=self._generate_element_id(),
                document_id=document_id,
                version_id=version_id,
                element_type=element_type,
                content=text,
                enhanced_content=text,
                reading_order=reading_order,
                confidence=0.95
            )
        except Exception:
            return None
