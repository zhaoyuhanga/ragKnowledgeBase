"""
RAG 问答系统 - 文件解析工具模块
支持 PDF、Markdown、TXT、DOCX 格式文件解析
"""

import hashlib
from pathlib import Path
from typing import Optional, List, Tuple
import re

import pypdf
import docx
import markdown

from app.core.logger import get_logger

logger = get_logger(__name__)


class FileParser:
    """
    文件解析器
    解析多种格式的文档文件，提取文本内容
    """
    
    @staticmethod
    def calculate_hash(file_path: str) -> str:
        """
        计算文件的 MD5 哈希值
        
        Args:
            file_path: 文件路径
            
        Returns:
            MD5 哈希值
        """
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    
    @staticmethod
    def parse_pdf(file_path: str) -> Tuple[str, int]:
        """
        解析 PDF 文件
        
        Args:
            file_path: PDF 文件路径
            
        Returns:
            (提取的文本, 页数)
        """
        logger.info(f"正在解析 PDF 文件: {file_path}")
        
        try:
            texts = []
            page_count = 0
            
            reader = pypdf.PdfReader(file_path)
            page_count = len(reader.pages)
            
            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        texts.append(text.strip())
                except Exception as e:
                    logger.warning(f"提取第 {page_num + 1} 页文本失败: {str(e)}")
            
            content = "\n\n".join(texts)
            logger.info(f"PDF 解析完成，页数: {page_count}, 字符数: {len(content)}")
            
            return content, page_count
            
        except Exception as e:
            logger.error(f"PDF 解析失败: {str(e)}")
            raise
    
    @staticmethod
    def parse_markdown(file_path: str) -> Tuple[str, int]:
        """
        解析 Markdown 文件
        
        Args:
            file_path: Markdown 文件路径
            
        Returns:
            (提取的文本, 行数)
        """
        logger.info(f"正在解析 Markdown 文件: {file_path}")
        
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            
            # 预处理 Markdown 文本
            # 移除一些常见的标记符号，保留主要内容
            lines = content.split("\n")
            processed_lines = []
            
            for line in lines:
                # 保留标题、列表、段落内容
                # 移除图片链接但保留 alt 文本
                line = re.sub(r'!\[.*?\]\(.*?\)', '', line)
                # 移除链接但保留文本
                line = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', line)
                # 移除代码块标记但保留内容
                line = line.replace('```', '')
                # 保留标题符号
                processed_lines.append(line)
            
            result = "\n".join(processed_lines)
            line_count = len(processed_lines)
            
            logger.info(f"Markdown 解析完成，行数: {line_count}, 字符数: {len(result)}")
            
            return result, line_count
            
        except Exception as e:
            logger.error(f"Markdown 解析失败: {str(e)}")
            raise
    
    @staticmethod
    def parse_txt(file_path: str) -> Tuple[str, int]:
        """
        解析 TXT 文件
        
        Args:
            file_path: TXT 文件路径
            
        Returns:
            (提取的文本, 行数)
        """
        logger.info(f"正在解析 TXT 文件: {file_path}")
        
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            
            line_count = len(content.split("\n"))
            
            logger.info(f"TXT 解析完成，行数: {line_count}, 字符数: {len(content)}")
            
            return content, line_count
            
        except Exception as e:
            logger.error(f"TXT 解析失败: {str(e)}")
            raise
    
    @classmethod
    def parse_docx(cls, file_path: str) -> Tuple[str, int]:
        """
        解析 DOCX 文件
        
        Args:
            file_path: DOCX 文件路径
            
        Returns:
            (提取的文本, 段落数)
        """
        logger.info(f"正在解析 DOCX 文件: {file_path}")
        
        try:
            from zipfile import BadZipFile
            doc = docx.Document(file_path)
            
            texts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    texts.append(para.text.strip())
            
            # 尝试提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            texts.append(cell.text.strip())
            
            content = "\n\n".join(texts)
            para_count = len(texts)
            
            logger.info(f"DOCX 解析完成，段落数: {para_count}, 字符数: {len(content)}")
            
            return content, para_count
            
        except BadZipFile:
            # 处理损坏的 DOCX 文件
            logger.warning(f"DOCX 文件格式异常，尝试使用备用方法: {file_path}")
            return cls._parse_docx_fallback(file_path)
        except Exception as e:
            if "NULL" in str(e) or "archive" in str(e).lower():
                # 处理嵌入对象导致的错误
                logger.warning(f"DOCX 包含嵌入对象，使用备用解析方法: {str(e)}")
                return cls._parse_docx_fallback(file_path)
            logger.error(f"DOCX 解析失败: {str(e)}")
            raise
    
    @classmethod
    def _parse_docx_fallback(cls, file_path: str) -> Tuple[str, int]:
        """备用 DOCX 解析方法，使用纯 ZIP 方式提取文本"""
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            
            texts = []
            with zipfile.ZipFile(file_path, 'r') as zf:
                for name in zf.namelist():
                    if name.startswith('word/document'):
                        with zf.open(name) as xml_file:
                            tree = ET.parse(xml_file)
                            root = tree.getroot()
                            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                            for elem in root.iter():
                                if elem.tag.endswith('}t') and elem.text:
                                    texts.append(elem.text)
            
            content = "\n\n".join(texts)
            para_count = len([t for t in texts if t.strip()])
            logger.info(f"DOCX 备用解析完成，段落数: {para_count}, 字符数: {len(content)}")
            return content, para_count
            
        except Exception as e:
            logger.error(f"DOCX 备用解析失败: {str(e)}")
            raise
    
    @classmethod
    def parse_file(cls, file_path: str, file_type: str) -> Tuple[str, int]:
        """
        根据文件类型解析文件
        
        Args:
            file_path: 文件路径
            file_type: 文件类型 (pdf/md/txt/docx)
            
        Returns:
            (提取的文本, 数量)
        """
        parsers = {
            "pdf": cls.parse_pdf,
            "md": cls.parse_markdown,
            "txt": cls.parse_txt,
            "docx": cls.parse_docx,
        }
        
        parser = parsers.get(file_type.lower())
        if not parser:
            raise ValueError(f"不支持的文件类型: {file_type}")
        
        return parser(file_path)
    
    @staticmethod
    def get_file_type(filename: str) -> Optional[str]:
        """
        从文件名获取文件类型
        
        Args:
            filename: 文件名
            
        Returns:
            文件类型（小写），不支持则返回 None
        """
        ext = Path(filename).suffix.lower().lstrip(".")
        
        type_mapping = {
            "pdf": "pdf",
            "md": "md",
            "markdown": "md",
            "txt": "txt",
            "docx": "docx",
        }
        
        return type_mapping.get(ext)


# 创建全局解析器实例
file_parser = FileParser()


def get_file_parser() -> FileParser:
    """获取文件解析器实例"""
    return file_parser
